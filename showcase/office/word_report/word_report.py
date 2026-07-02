"""
[Showcase / Office / 通用]
Skill: 从结构化内容生成 Word 文档

适用场景：
  通知、简报这类格式固定的文档，标题 + 分段正文，内容每次变、
  格式不变，直接用 python-docx 生成，不需要打开 Word，不需要屏幕。

用法：
  python run.py showcase/office/word_report/word_report -- \\
    --output report.docx \\
    --title "本周数据简报" \\
    --data '["第一段内容……", "第二段内容……"]'
"""

import argparse
import json
import sys

from docx import Document

from core.logger import SkillLogger


def generate_docx(title: str, paragraphs: list[str], output: str) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(output)


def main():
    parser = argparse.ArgumentParser(description="从结构化内容生成 Word 文档")
    parser.add_argument("--title", required=True, help="文档标题")
    parser.add_argument("--data", required=True, help="正文段落 JSON 数组（字符串列表）")
    parser.add_argument("--output", required=True, help="输出的 .docx 路径")
    try:
        sep = sys.argv.index("--")
        argv = sys.argv[sep + 1:]
    except ValueError:
        argv = sys.argv[1:]
    args = parser.parse_args(argv)

    log = SkillLogger("office/word_report")

    paragraphs = json.loads(args.data)
    log.step(f"生成文档：{args.title}（{len(paragraphs)} 段）")
    generate_docx(args.title, paragraphs, args.output)
    print(f"已保存到 {args.output}")
    log.finish({"paragraphs": len(paragraphs), "output": args.output})
    return {"output": args.output}


if __name__ == "__main__":
    main()
